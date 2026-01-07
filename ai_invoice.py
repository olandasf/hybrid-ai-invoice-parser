import os
import json
import logging
import re
import time
import openai

from typing import Any, Dict, List, Union, Optional

from google.cloud import documentai_v1 as documentai
from google.oauth2 import service_account

from akcizai import CATEGORY_LABELS, TARIFAI
from category import classify_alcohol
from simple_cache import SimpleFileCache, pdf_cache
from utils import clean_and_convert_to_float, clean_volume_value, log_function_call, clean_product_name
from image_preprocessing import preprocess_pdf

# --- FAILŲ TIPŲ APDOROJIMAS ---
def get_mime_type(file_path: str) -> str:
    """
    Nustato failo MIME tipą pagal plėtinį.
    Document AI palaiko: application/pdf, image/png, image/jpeg, image/tiff, image/gif, image/webp
    """
    ext = file_path.lower().rsplit('.', 1)[-1] if '.' in file_path else ''
    mime_types = {
        'pdf': 'application/pdf',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'tiff': 'image/tiff',
        'tif': 'image/tiff',
        'gif': 'image/gif',
        'webp': 'image/webp',
    }
    return mime_types.get(ext, 'application/pdf')

def convert_word_to_pdf(file_path: str) -> Optional[str]:
    """
    Konvertuoja Word dokumentą į PDF naudojant python-docx ir reportlab arba docx2pdf.
    Grąžina PDF failo kelią arba None jei konvertavimas nepavyko.
    """
    import tempfile
    output_pdf = file_path.rsplit('.', 1)[0] + '_converted.pdf'
    
    try:
        # Bandome naudoti docx2pdf (Windows)
        from docx2pdf import convert
        convert(file_path, output_pdf)
        logging.info(f"Word dokumentas konvertuotas į PDF: {output_pdf}")
        return output_pdf
    except ImportError:
        logging.warning("docx2pdf biblioteka nerasta, bandoma alternatyva...")
    except Exception as e:
        logging.warning(f"docx2pdf klaida: {e}")
    
    try:
        # Alternatyva: naudojame python-docx ir reportlab
        from docx import Document as DocxDocument
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        
        doc = DocxDocument(file_path)
        c = canvas.Canvas(output_pdf, pagesize=A4)
        width, height = A4
        
        y_position = height - 2*cm
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                # Paprastas teksto įterpimas (be formatavimo)
                if y_position < 2*cm:
                    c.showPage()
                    y_position = height - 2*cm
                c.drawString(2*cm, y_position, text[:100])  # Limituojame eilutės ilgį
                y_position -= 0.5*cm
        
        c.save()
        logging.info(f"Word dokumentas konvertuotas į PDF (alternatyva): {output_pdf}")
        return output_pdf
    except ImportError as e:
        logging.error(f"Trūksta bibliotekų Word konvertavimui: {e}")
    except Exception as e:
        logging.error(f"Klaida konvertuojant Word į PDF: {e}")
    
    return None

def is_word_document(file_path: str) -> bool:
    """Tikrina ar failas yra Word dokumentas."""
    ext = file_path.lower().rsplit('.', 1)[-1] if '.' in file_path else ''
    return ext in ['doc', 'docx']

# --- KONFIGŪRACIJA ---
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"

deepseek_client = None
if DEEPSEEK_API_KEY and DEEPSEEK_API_KEY not in ["ĮRAŠYKITE_SAVO_DEEPSEEK_API_RAKTĄ_ČIA", "your-deepseek-api-key-here"]:
    try:
        deepseek_client = openai.OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
        logging.info("Sukurtas DeepSeek API klientas.")
    except Exception as e:
        logging.error(f"Nepavyko sukurti DeepSeek API kliento: {e}", exc_info=True)
else:
    logging.warning(f"DeepSeek API raktas nenurodytas arba neteisingas: '{DEEPSEEK_API_KEY}'. DeepSeek funkcijos neveiks.")

# Naujas cache produktų klasifikavimui
product_classification_cache = SimpleFileCache(cache_dir="product_cache", max_age_hours=720) # 30 dienų

# --- KONSTANTOS ---
PACKAGING_KEYWORDS = [
    'olive oil', 'olio ', 'olio di', ' aliejus', 'aceite', 'extra virgin', 'alyvuogių', 'olim', 'vidre', 
    'pallet', 'palet', 
    'empty box', 'empty gift box', 'tuščia dėžutė', 'gift box', 'giftbox', 'single gift box',
    'wooden box', 'wood box', 'packaging', 'pakuotė'
]

NON_ALCOHOLIC_DRINK_KEYWORDS = [
    'alcohol free', 'alcohol-free', 'alcoholvrij', 'alkoholfrei', 'sans alcool', 'senza alcol', 'sin alcohol', 'nealkoholinis', 'no alcohol'
]

NON_ALCOHOLIC_KEYWORDS = PACKAGING_KEYWORDS + NON_ALCOHOLIC_DRINK_KEYWORDS

# --- PAGALBINĖS FUNKCIJOS ---

def classify_product_with_deepseek(product_name: str, abv: float) -> str:
    """
    Klasifikuoja produktą naudojant DeepSeek AI modelį.
    Naudoja cache, kad išvengtų pasikartojančių API užklausų.
    """
    logging.info(f"classify_product_with_deepseek iškviesta: product_name='{product_name}', abv={abv}")
    if not product_name or not deepseek_client:
        logging.warning("Produkto pavadinimas tuščias arba DeepSeek klientas nepasiekiamas. Naudojamas fallback.")
        return classify_alcohol(product_name, abv)

    # 0. Priverstinis patikrinimas dėl specifinių vynų (PRIEŠ CACHE)
    # Tai užtikrina, kad net jei cache yra klaidingas įrašas, šie produktai bus klasifikuojami teisingai
    forced_wine_keywords = ['acediano']
    if any(k in product_name.lower() for k in forced_wine_keywords):
        logging.info(f"Produktas '{product_name}' priverstinai klasifikuojamas kaip vynas (apeinant Cache ir DeepSeek).")
        if abv > 15.0:
             return 'wine_8.5_15' 
        elif abv > 8.5:
             return 'wine_8.5_15'
        else:
             return 'wine_up_to_8.5'

    # 0.1 Priverstinis stipriųjų gėrimų patikrinimas (PRIEŠ CACHE)
    forced_spirits_keywords = ['navimer', 'alcohol pur', 'rectified spirit', 'spirytus']
    if any(k in product_name.lower() for k in forced_spirits_keywords):
        logging.info(f"Produktas '{product_name}' priverstinai klasifikuojamas kaip etilo alkoholis.")
        return 'ethyl_alcohol'

    # 0.2 Priverstinis nealkoholinių produktų patikrinimas (PRIEŠ CACHE)
    if any(k in product_name.lower() for k in NON_ALCOHOLIC_KEYWORDS):
        logging.info(f"Produktas '{product_name}' priverstinai klasifikuojamas kaip nealkoholinis (pagal raktažodžius).")
        return 'non_alcohol'

    cache_key = f"{product_name}_{abv}"
    cached_data = product_classification_cache.get(cache_key)
    if cached_data and 'category' in cached_data:
        logging.info(f"DeepSeek klasifikacija rasta cache: '{product_name}' -> {cached_data['category']}")
        return cached_data['category']
    else:
        logging.info(f"DeepSeek klasifikacija NĖRA rasta cache raktui: '{cache_key}'")

    # Sukuriame kategorijų sąrašą paaiškinimui
    categories_text = "\n".join([f"- `{key}`: {description}" for key, description in CATEGORY_LABELS.items()])

    prompt_content = f"""
Tu esi pasaulinio lygio akcizo mokesčio ekspertas, specializuojantis alkoholinių gėrimų klasifikavime. Tavo užduotis yra **absoliučiai tiksliai** klasifikuoti gėrimą pagal pavadinimą ir ABV.

**Produkto pavadinimas:** "{product_name}"
**Alkoholio stiprumas (ABV):** {abv}%

**Galimos kategorijos:**
{categories_text}

**GELEŽINĖS TAISYKLĖS (CRITICAL RULES) - PRIVALOMA LAIKYTIS:**
1.  **STIPRIEJI GĖRIMAI (SPIRITS):** Viskis (Whisky/Whiskey), Romas (Rum), Konjakas (Cognac), Brendis (Brandy), Tekila (Tequila), Džinas (Gin), Degtinė (Vodka), Arakas (Arak), Grappa, Pisco, Calvados yra **VISADA** `ethyl_alcohol`. Tai svarbiausia taisyklė.
2.  **LIKERIAI (LIQUEURS):** Visi likeriai, įskaitant, bet neapsiribojant, 'St Germain Elderflower', 'Sheridan's', 'Cointreau', 'Jägermeister', 'Amaretto', 'Baileys' yra **VISADA** `ethyl_alcohol`.
3.  **AMARO / BITTERS:** Gėrimai kaip 'Fernet Branca', 'Branca Menta', 'Aperol', 'Campari' yra **VISADA** `ethyl_alcohol`.
4.  **IGNORUOTI PRIEDUS:** Visiškai ignoruok priedus pavadinime, tokius kaip '+ GB', '+ Glass', 'Gift Box', 'dėžutė', 'su stikline'. Jie nekeičia gėrimo esmės. "Glen Grant 12 Years + GB" yra viskis, todėl `ethyl_alcohol`.
5.  **PUTOJANTYS VYNAI:** Šampanas (Champagne), Prosecco, Cava, Spumante, Asti klasifikuojami pagal ABV į `sparkling_wine_le_8_5` arba `sparkling_wine_over_8_5`.
6.  **STIPRINTI VYNAI:** Portas (Port), Cheresas (Sherry), Madeira, Marsala klasifikuojami kaip `intermediate_products_...` pagal jų ABV.
7.  **TYLŪS VYNAI:** Visi kiti vynai (Wine, Vin, Vino) klasifikuojami pagal ABV į atitinkamas vyno kategorijas.
8.  **NE GĖRIMAI:** Produktai, kurie nėra gėrimai (pvz., 'Glass', 'Gift Box', 'Dėžutė', 'Taurė'), yra `non_alcohol`.

**PROCESAS IR PAVYZDŽIAI:**
1.  Perskaityk pavadinimą ir ABV.
2.  Pritaikyk **GELEŽINES TAISYKLES**.
3.  Grąžink **TIK vieną** kategorijos raktą. Jokio papildomo teksto.

*   **Pavyzdys 1 (Viskis su dovanų dėžute):**
    *   Pavadinimas: "The Glendronach 12 Years + GB"
    *   Taisyklė: Nr. 1 (Stiprieji gėrimai) ir Nr. 4 (Ignoruoti priedus).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 2 (Romas):**
    *   Pavadinimas: "Barcelo Imperial Onyx + GB"
    *   Taisyklė: Nr. 1 (Stiprieji gėrimai).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 3 (Likeris):**
    *   Pavadinimas: "Sheridan's Double"
    *   Taisyklė: Nr. 2 (Likeriai).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 4 (Bitter):**
    *   Pavadinimas: "Fernet Branca"
    *   Taisyklė: Nr. 3 (Amaro / Bitters).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 5 (Tekila):**
    *   Pavadinimas: "Corazon Blanco"
    *   Taisyklė: Nr. 1 (Stiprieji gėrimai).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 6 (Konjakas):**
    *   Pavadinimas: "Frapin VIP XO + GB"
    *   Taisyklė: Nr. 1 (Stiprieji gėrimai).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 7 (Viskis):**
    *   Pavadinimas: "Tomatin Legacy + GB"
    *   Taisyklė: Nr. 1 (Stiprieji gėrimai) ir Nr. 4 (Ignoruoti priedus).
    *   Rezultatas: `ethyl_alcohol`

*   **Pavyzdys 8 (Likeris):**
    *   Pavadinimas: "St Germain Elderflower"
    *   Taisyklė: Nr. 2 (Likeriai).
    *   Rezultatas: `ethyl_alcohol`
    
*   **Pavyzdys 9 (Romas):**
    *   Pavadinimas: "Barcelo Imperial Premium Blend 40th Anniversary"
    *   Taisyklė: Nr. 1 (Stiprieji gėrimai).
    *   Rezultatas: `ethyl_alcohol`

**TAVO UŽDUOTIS DABAR:**
Pateik **TIK KATEGORIJOS RAKTĄ** žemiau nurodytam produktui. Griežtai laikykitės taisyklių. Bet koks nukrypimas nuo taisyklių yra nepriimtinas.
"""

    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "Tu esi pasaulinio lygio akcizo mokesčio ekspertas. Tavo vienintelė užduotis - klasifikuoti gėrimą pagal pateiktas taisykles ir grąžinti TIK kategorijos raktą. Jokių paaiškinimų, jokio papildomo teksto."},
            {"role": "user", "content": prompt_content}
        ],
        "max_tokens": 50,
        "temperature": 0.0
    }

    try:
        logging.info(f"Kreipiamasi į DeepSeek dėl produkto klasifikavimo: '{product_name}' (ABV: {abv})")
        # Pridėtas timeout=10s, kad neužstrigtų
        response = deepseek_client.chat.completions.create(**payload, timeout=10)
        content = response.choices[0].message.content
        logging.info(f"DeepSeek atsakymas (raw): '{content}'")
        category_key = content.strip().replace('`', '')

        if category_key in CATEGORY_LABELS:
            logging.info(f"DeepSeek sėkmingai klasifikavo '{product_name}' kaip '{category_key}'")
            product_classification_cache.set(cache_key, {'category': category_key})
            return category_key
        else:
            logging.warning(f"DeepSeek grąžino nežinomą kategoriją: '{category_key}'. Naudojamas fallback.")
            fallback_category = classify_alcohol(product_name, abv)
            logging.info(f"Fallback klasifikacija: '{product_name}' -> '{fallback_category}'")
            return fallback_category

    except Exception as e:
        logging.error(f"DeepSeek API klaida klasifikuojant produktą: {e}. Naudojamas fallback.", exc_info=True)
        fallback_category = classify_alcohol(product_name, abv)
        logging.info(f"Fallback klasifikacija po klaidos: '{product_name}' -> '{fallback_category}'")
        return fallback_category


def validate_transport_amount(amount: float) -> tuple[bool, str]:
    """
    Validuoja transporto sumą.
    
    Args:
        amount: Transporto suma eurais
        
    Returns:
        tuple: (ar_validus, klaidos_pranešimas)
    """
    if amount < 0:
        return False, "Transporto suma negali būti neigiama"
    if amount > 10000:
        return False, "Transporto suma per didelė (>10000 EUR)"
    return True, ""

def extract_transport_from_document(document) -> float:
    """
    Ištraukia transporto sumą iš Document AI response.
    
    Args:
        document: Document AI response objektas
        
    Returns:
        float: Transporto suma EUR arba 0.0 jei nerastas
    """
    try:
        transport_total = 0.0
        
        # Ieškome transporto eilučių Document AI entities
        if hasattr(document, 'entities'):
            for entity in document.entities:
                # Tikriname ar tai transporto eilutė
                if hasattr(entity, 'type_') and 'transport' in entity.type_.lower():
                    # Bandome ištraukti sumą
                    if hasattr(entity, 'normalized_value') and hasattr(entity.normalized_value, 'money_value'):
                        amount = entity.normalized_value.money_value.units
                        if amount > 0:
                            transport_total += amount
                            logging.info("Document AI rado transporto eilutę: %.2f EUR", amount)
        
        # Jei entities neveikia, bandome per line_items
        if transport_total == 0.0 and hasattr(document, 'pages'):
            for page in document.pages:
                if hasattr(page, 'tables'):
                    for table in page.tables:
                        if hasattr(table, 'body_rows'):
                            for row in table.body_rows:
                                if hasattr(row, 'cells'):
                                    row_text = ""
                                    amount_text = ""
                                    for cell in row.cells:
                                        if hasattr(cell, 'layout') and hasattr(cell.layout, 'text_anchor') and cell.layout.text_anchor.text_segments:
                                            cell_text = document.text[cell.layout.text_anchor.text_segments[0].start_index:
                                                                   cell.layout.text_anchor.text_segments[0].end_index]
                                            # Tikriname ar tai transporto raktažodis
                                            has_transport_keyword = any(keyword in cell_text.lower() for keyword in 
                                                  ['freight', 'transport', 'fracht', 'livraison', 'spedizione', 'vracht', 'transporte'])
                                            
                                            # Tikriname ar tai nėra produktas
                                            is_product = any(keyword in cell_text.lower() for keyword in 
                                                  ['pallets', 'pallet', 'jack daniel', 'whiskey', 'whisky', 'vodka', 'gin', 'rum', 'wine', 'vynas', 'champagne', 'prosecco', 'beer', 'alus'])
                                            
                                            if has_transport_keyword and not is_product:
                                                row_text = cell_text
                                            # Bandome rasti sumą toje pačioje eilutėje
                                            amount_match = re.search(r'(\d+[.,]\d+|\d+)', cell_text)
                                            if amount_match and row_text:
                                                amount_text = amount_match.group(1)
                                    
                                    if row_text and amount_text:
                                        amount = clean_and_convert_to_float(amount_text)
                                        if amount and amount > 0:
                                            transport_total += amount
                                            logging.info("Document AI rado transporto eilutę lentelėje: %s - %.2f EUR", row_text, amount)
        
        # Validuojame rastą sumą
        is_valid, error_msg = validate_transport_amount(transport_total)
        if not is_valid:
            logging.warning(f"Transporto validacijos klaida: {error_msg}")
            return 0.0
            
        if transport_total > 0:
            logging.info("Bendras rastas transportas: %.2f EUR", transport_total)
        
        return transport_total
        
    except Exception as e:
        logging.error(f"Klaida ištraukiant transportą iš Document AI: {e}", exc_info=True)
        return 0.0

def filter_transport_lines(products: list) -> tuple[list, float]:
    """
    Atskiria transporto eilutes nuo produktų.
    
    Args:
        products: Produktų sąrašas iš AI ekstraktavimo
        
    Returns:
        tuple: (išfiltruoti_produktai, papildoma_transporto_suma)
    """
    # Transporto raktažodžiai skirtingomis kalbomis
    transport_keywords = [
        # Lietuvių
        'transportas', 'gabenimas', 'pristatymas', 'vežimas',
        # Anglų
        'freight', 'transport', 'shipping', 'delivery', 'carriage',
        # Vokiečių
        'fracht', 'versand', 'lieferung',
        # Olandų
        'vracht', 'verzending', 'levering',
        # Prancūzų
        'fret', 'livraison', 'expédition',
        # Ispanų
        'flete', 'envío', 'entrega',
        # Italų
        'spedizione', 'consegna', 'nolo'
    ]
    
    # Produktai, kurie NIEKADA nėra transportas (net jei turi transporto raktažodžius)
    non_transport_keywords = [
        'jack daniel', 'whiskey', 'whisky', 'vodka', 'gin', 'rum',
        'wine', 'vynas', 'champagne', 'prosecco', 'beer', 'alus'
    ]

    # Eilutės, kurios apskritai neturėtų būti laikomos produktais (pvz. paletės)
    pallet_keywords = [
        'pallet', 'pallets', 'palete', 'paletė', 'paletės',
        'fumigated pallet', 'euro pallet', 'euro-pallet', 'epal', 'ippc marked',
        'palet', 'consolidation fees'
    ]
    
    filtered_products = []
    transport_amount = 0.0
    
    for product in products:
        product_name = str(product.get('name', '')).lower()

        # 0. Paletės – visiškai pašaliname iš produktų ir nelaikome nei transportu, nei preke
        # TAČIAU: Jei eilutė turi ir produkto požymių (pvz. Underberg), tai tik išvalome žodį "Pallet"
        if any(keyword in product_name for keyword in pallet_keywords):
            # Tikriname ar tai nėra sumaišyta eilutė (Produktas + Paletė)
            # Ieškome stiprių produkto indikatorių
            # Papildyti raktažodžiai: spiritai, metai, pakuotės tipai
            mixed_indicators = [
                'underberg', 'bottles', 'liter', '%', 'vol', 'cl', 'ml', 
                'whisky', 'whiskey', 'vodka', 'gin', 'rum', 'cognac', 'brandy',
                'years', 'yr', 'old', 'aged', 'carton', 'case', 'gb'
            ]
            is_mixed_product = any(k in product_name for k in mixed_indicators)
            
            # Taip pat tikriname ar produktas turi validžius kiekius/kainas/tūrį/ABV
            if not is_mixed_product:
                has_valid_data = (
                    (product.get('quantity', 0) or 0) > 0 or 
                    (product.get('unit_price', 0) or 0) > 5.0 or  # Paletės paprastai pigios arba neturi vnt. kainos
                    (product.get('volume', 0) or 0) > 0 or
                    (product.get('abv', 0) or 0) > 0
                )
                if has_valid_data:
                    is_mixed_product = True

            if is_mixed_product:
                logging.info(f"Eilutė turi paletės raktažodžių, bet atrodo kaip produktas. Valomas pavadinimas: {product.get('name')}")
                # Išvalome paletės raktažodžius iš pavadinimo
                clean_name = product.get('name', '')
                for pk in pallet_keywords:
                    # Case-insensitive replace
                    pattern = re.compile(re.escape(pk), re.IGNORECASE)
                    clean_name = pattern.sub('', clean_name)
                product['name'] = clean_name.strip()
                # Toliau apdorojame kaip normalų produktą
                product_name = product['name'].lower()
            else:
                logging.info(f"Eilutė laikoma palete ir pašalinama iš produktų: {product.get('name')}")
                continue
        
        # Pirmiausia tikriname ar tai tikrai NĖRA produktas
        is_definitely_product = any(keyword in product_name for keyword in non_transport_keywords)
        
        # Tikriname ar produkto pavadinime yra transporto raktažodžių
        has_transport_keywords = any(keyword in product_name for keyword in transport_keywords)
        
        # Tai transportas tik jei turi transporto raktažodžių IR nėra tikras produktas
        is_transport = has_transport_keywords and not is_definitely_product
        
        if is_transport:
            # Tai transporto eilutė - ištraukiame sumą
            amount = clean_and_convert_to_float(product.get('amount', 0))
            if amount and amount > 0:
                transport_amount += amount
                logging.info(f"Rasta transporto eilutė produktuose: {product.get('name')} - {amount} EUR")
        else:
            # Tai tikras produktas - paliekame sąraše
            filtered_products.append(product)
    
    return filtered_products, transport_amount

def safe_extract_transport(document) -> float:
    """
    Saugus transporto ekstraktavimas su klaidų apdorojimu.
    
    Args:
        document: Document AI response objektas
        
    Returns:
        float: Transporto suma EUR arba 0.0 jei klaida
    """
    if not document:
        logging.warning("Document AI response yra None - transporto aptikimas praleistas")
        return 0.0
        
    try:
        transport_amount = extract_transport_from_document(document)
        if transport_amount > 0:
            logging.info(f"Sėkmingai aptiktas transportas Document AI: {transport_amount} EUR")
        else:
            logging.debug("Document AI transporto nerado")
        return transport_amount
    except AttributeError as e:
        logging.warning(f"Document AI struktūros klaida transporto aptikime: {e}")
        return 0.0
    except ValueError as e:
        logging.warning(f"Transporto sumos konvertavimo klaida: {e}")
        return 0.0
    except Exception as e:
        logging.error(f"Netikėta transporto aptikimo klaida: {e}", exc_info=True)
        return 0.0

def estimate_abv_from_name(product_name: str) -> float:
    """Nustato tipinį alkoholio stiprumą pagal produkto pavadinimą"""
    if not product_name:
        return 0.0
    
    name_lower = product_name.lower()
    
    # 1. Pirmiausia ieškome explicitiško ABV nurodymo pavadinime (pvz. "38%", "40 %", "12.5%")
    # Ieškome skaičiaus prieš '%' ženklą.
    abv_match = re.search(r'(\d+[.,]?\d*)\s*%', name_lower)
    if abv_match:
        try:
            abv_val = float(abv_match.group(1).replace(',', '.'))
            # Validacija: ABV turi būti protingose ribose (pvz. 0.5 - 99)
            if 0.5 <= abv_val <= 99.0:
                logging.info(f"ABV rasta pavadinime '{product_name}': {abv_val}%")
                return abv_val
        except ValueError:
            pass

    # Jei produktas neturi pavadinimo (našlaitė), jis negali būti alkoholis
    if '(be pavadinimo)' in name_lower:
        return 0.0
    
    # Nealkoholiniai produktai (Aliejus, Maistas ir pan.)
    if any(keyword in name_lower for keyword in NON_ALCOHOLIC_KEYWORDS):
        return 0.0
    
    # Stiprūs alkoholiniai gėrimai (40%)
    spirits_keywords = [
        'vodka', 'whisky', 'whiskey', 'gin', 'rum', 'ron', 'cognac', 'brandy', 'tequila', 'bourbon', 'scotch', 'jack daniel',
        'ardbeg', 'auchentoshan', 'balvenie', 'connemara', 'dalwhinnie', 'finlaggan', 'glenallachie', 'glendronach', 'glenfiddich',
        'grand marnier', 'jura', 'zacapa', 'unicum', 'underberg', 'deanston', 'arran', 'talisker', 'laphroaig', 'lagavulin',
        'bowmore', 'bunnahabhain', 'caol ila', 'bruichladdich', 'kilchoman', 'highland park', 'macallan', 'glenmorangie',
        'glenlivet', 'aberlour', 'balblair', 'benriach', 'benromach', 'bladnoch', 'dalmore', 'fettercairn', 'glengoyne',
        'glenrothes', 'knockando', 'mortlach', 'old pulteney', 'royal lochnagar', 'springbank', 'tullibardine', 'wolfburn',
        'aberfeldy', 'aberalour', 'ardmore', 'arran', 'auchentoshan', 'auld', 'balblair', 'balvenie', 'ben nevis', 'benriach',
        'benrinnes', 'benromach', 'blair athol', 'bowmore', 'bruichladdich', 'bunnahabhain', 'caol ila', 'cardhu', 'clynelish',
        'cragganmore', 'dailuaine', 'dalmore', 'dalwhinnie', 'deanston', 'edradour', 'fettercairn', 'glen elgin', 'glen garioch',
        'glen grant', 'glen keith', 'glen moray', 'glen ord', 'glen scotia', 'glen spey', 'glenallachie', 'glenburgie',
        'glencadam', 'glendronach', 'glendullan', 'glenfarclas', 'glenfiddich', 'glenglassaugh', 'glengoyne', 'glenkinchie',
        'glenlivet', 'glenlossie', 'glenmorangie', 'glenrothes', 'glenturret', 'highland park', 'inchgower', 'jura',
        'kilchoman', 'knockando', 'lagavulin', 'laphroaig', 'linkwood', 'loch lomond', 'longmorn', 'macallan', 'mannochmore',
        'miltonduff', 'mortlach', 'oban', 'old pulteney', 'royal brackla', 'royal lochnagar', 'scapa', 'speyburn', 'springbank',
        'strathisla', 'strathmill', 'talisker', 'tamdhu', 'tamnavulin', 'teaninich', 'tobermory', 'tomatin', 'tomintoul',
        'tormore', 'tullibardine', 'wolfburn',
        'fernet', 'branca', 'amaro', 'campari', 'aperol', 'jagermeister', 'underberg'
    ]
    if any(keyword in name_lower for keyword in spirits_keywords):
        return 40.0
    
    # Likieriai (25%)
    liqueur_keywords = ['liqueur', 'likeris', 'likor', 'amaretto', 'baileys', 'kahlua', 'sambuca']
    if any(keyword in name_lower for keyword in liqueur_keywords):
        return 25.0
    
    # Tarpiniai produktai
    intermediate_keywords = ['port', 'porto', 'sherry', 'madeira', 'marsala']
    if any(keyword in name_lower for keyword in intermediate_keywords):
        return 20.0
    
    # Vermutas
    if 'vermouth' in name_lower or 'vermutas' in name_lower:
        return 16.0
    
    # Aukšto ABV vynai
    high_wine_keywords = ['amarone', 'ripasso', 'primitivo di manduria']
    if any(keyword in name_lower for keyword in high_wine_keywords):
        return 15.5
    
    # Italų premium vynai
    premium_wine_keywords = ['barolo', 'barbaresco', 'brunello']
    if any(keyword in name_lower for keyword in premium_wine_keywords):
        return 14.0
    
    # Chianti (paprastas vynas)
    if 'chianti' in name_lower:
        return 13.0
    
    # Putojantys vynai
    sparkling_keywords = ['champagne', 'prosecco', 'cava', 'sparkling', 'spumante', 'franciacorta', 'brut']
    if any(keyword in name_lower for keyword in sparkling_keywords):
        return 12.0
    
    # Vynai
    wine_keywords = ['wine', 'vin', 'vino', 'wein', 'rouge', 'blanc', 'rosso', 'bianco', 'red', 'white', 'rose']
    if any(keyword in name_lower for keyword in wine_keywords):
        return 13.0
    
    # Alus
    beer_keywords = ['beer', 'bier', 'birra', 'cerveza', 'ale', 'lager', 'stout', 'ipa']
    if any(keyword in name_lower for keyword in beer_keywords):
        return 5.0
    
    # Jei negalime nustatyti, bet produktas atrodo kaip alkoholinis
    # (turi tūrį litrais ir nėra aiškiai nealkoholinis)
    return 12.0  # Tipinis vyno stiprumas kaip default

# clean_and_convert_to_float funkcija perkelta į utils.py


def _get_entity_value(entity: Any) -> Union[str, float]:
    """Grąžina normalizuotą reikšmę iš Document AI sub-entito."""
    try:
        normalized = getattr(entity, 'normalized_value', None)
        if normalized:
            money_value = getattr(normalized, 'money_value', None)
            if money_value:
                units = getattr(money_value, 'units', 0) or 0
                nanos = getattr(money_value, 'nanos', 0) or 0
                return float(units) + float(nanos) / 1_000_000_000
            percent_value = getattr(normalized, 'percent_value', None)
            if percent_value is not None:
                return float(percent_value)
            text_value = getattr(normalized, 'text', None)
            if text_value and text_value.strip():
                return text_value.strip()
        mention = getattr(entity, 'mention_text', None)
        if mention and mention.strip():
            return mention.strip()
    except Exception as error:
        logging.debug("Nepavyko ištraukti normalizuotos reikšmės: %s", error)
    return ""


def extract_text_at_row(document, page_index, y_center, tolerance=0.02) -> str:
    """Ištraukia visą tekstą iš nurodytos eilutės (pagal Y koordinatę)."""
    try:
        if page_index >= len(document.pages):
            return ""
            
        page = document.pages[page_index]
        # Naudojame lines, nes jos dažniausiai geriau sugrupuoja tekstą
        items = page.lines if page.lines else page.tokens
        
        candidates = []
        
        for item in items:
            try:
                vertices = item.layout.bounding_poly.normalized_vertices
                y_item = sum(v.y for v in vertices) / len(vertices)
                
                # Use a slightly larger tolerance to collect candidates
                if abs(y_item - y_center) < tolerance:
                    text = ""
                    for segment in item.layout.text_anchor.text_segments:
                        start = int(segment.start_index)
                        end = int(segment.end_index)
                        text += document.text[start:end]
                    
                    if text.strip():
                        # Store X coordinate for sorting
                        x_coord = vertices[0].x if vertices else 0.0
                        candidates.append({
                            'text': text.strip(),
                            'y': y_item,
                            'x': x_coord,
                            'diff': abs(y_item - y_center)
                        })
            except:
                continue
        
        if not candidates:
            return ""

        # Group candidates by Y coordinate (clustering)
        # We assume lines are separated by at least 0.01
        clusters = []
        candidates.sort(key=lambda c: c['y'])
        
        current_cluster = [candidates[0]]
        for i in range(1, len(candidates)):
            if candidates[i]['y'] - candidates[i-1]['y'] < 0.01:
                current_cluster.append(candidates[i])
            else:
                clusters.append(current_cluster)
                current_cluster = [candidates[i]]
        clusters.append(current_cluster)
        
        # Find the cluster closest to y_center
        best_cluster = None
        min_cluster_diff = float('inf')
        
        for cluster in clusters:
            # Calculate average Y for the cluster
            avg_y = sum(c['y'] for c in cluster) / len(cluster)
            diff = abs(avg_y - y_center)
            
            if diff < min_cluster_diff:
                min_cluster_diff = diff
                best_cluster = cluster
        
        # Sort items in the best cluster by X coordinate
        best_cluster.sort(key=lambda c: c['x'])
        
        result = " ".join([c['text'] for c in best_cluster])
        logging.info(f"Recovered text at Y={y_center:.4f}: '{result}'")
        return result
        
    except Exception as e:
        logging.warning(f"Klaida traukiant tekstą eilutėje: {e}")
        return ""

def is_pallet_or_packaging_row(document, page_index, y_center, tolerance=0.02):
    """Patikrina ar nurodytoje eilutėje yra tekstas, nurodantis paletes ar pakuotes."""
    text = extract_text_at_row(document, page_index, y_center, tolerance)
    if not text:
        return False
        
    text_lower = text.lower()
    keywords = ['pallet', 'palete', 'paletė', 'eur-epal', 'europallet', 'padėklas']
    if any(k in text_lower for k in keywords):
        logging.info(f"Rasta paletė eilutėje (Y={y_center:.4f}): '{text.strip()}'")
        return True
    return False


def extract_line_items_from_doc_ai(document) -> list:
    """
    Ištraukia prekių eilutes iš Document AI response.
    Naudoja UNIFIKUOTĄ metodą: surenka visus entitus (tiek plokščius, tiek iš 'line_item' struktūrų)
    ir grupuoja juos pagal Y koordinates. Tai išsprendžia problemą, kai Document AI randa tik dalį 'line_item' eilučių.
    """
    logging.info("Pradedamas prekių eilučių ištraukimas (UNIFIKUOTAS METODAS)...")
    
    # --- Paletės filtravimas ---
    pallet_entities = [e for e in document.entities if e.type_ in ['pallet_line', 'packaging_fee']]
    if pallet_entities:
        logging.info(f"Rasta {len(pallet_entities)} 'pallet_line' entitų. Susiję duomenys bus ignoruojami.")

    def is_part_of_pallet(entity):
        if not pallet_entities: return False
        try:
            if not entity.text_anchor or not entity.text_anchor.text_segments: return False
            e_start = int(entity.text_anchor.text_segments[0].start_index)
            e_end = int(entity.text_anchor.text_segments[0].end_index)
            for pallet in pallet_entities:
                if not pallet.text_anchor or not pallet.text_anchor.text_segments: continue
                for p_seg in pallet.text_anchor.text_segments:
                    p_start = int(p_seg.start_index)
                    p_end = int(p_seg.end_index)
                    if max(e_start, p_start) < min(e_end, p_end):
                        return True
        except: pass
        return False
    # ---------------------------

    # 1. Surenkame visus entitus (Flattening)
    flat_entities = []
    
    # Pridedame top-level entitus
    for entity in document.entities:
        if entity.type_ == 'line_item':
            # Išpakuojame line_item savybes
            if entity.properties:
                for prop in entity.properties:
                    flat_entities.append(prop)
        else:
            flat_entities.append(entity)
            
    # 2. Filtruojame aktualius tipus
    # FIX: Added lowercase 'volume' since Google may return it with different casing
    relevant_types = [
        'product_name', 'product_code', 'quantity', 'unit_price', 'amount', 'Volume', 'volume', 'abv', 'description',
        'line_item/description', 'line_item/product_code', 'line_item/quantity', 'line_item/unit_price', 
        'line_item/amount', 'line_item/volume', 'line_item/abv'
    ]
    
    # DEBUG: Log all entity types returned by Google
    all_entity_types = set()
    for entity in document.entities:
        all_entity_types.add(entity.type_)
        if entity.type_ == 'line_item' and entity.properties:
            for prop in entity.properties:
                all_entity_types.add(prop.type_)
    logging.info(f"DEBUG: All entity types from Google: {sorted(all_entity_types)}")
    
    all_entities = [
        e for e in flat_entities 
        if e.type_ in relevant_types 
        and e.page_anchor.page_refs 
        and not is_part_of_pallet(e)
    ]
    
    logging.info(f"Iš viso rasta {len(all_entities)} aktualių entitų (po išpakavimo ir filtravimo).")

    # 3. Grupuojame pagal Y koordinates (Strategy 2 logika)
    
    def get_sort_key(e):
        try:
            page_ref = e.page_anchor.page_refs[0]
            page_index = int(page_ref.page) if page_ref.page else 0
            vertices = page_ref.bounding_poly.normalized_vertices
            y_avg = sum(v.y for v in vertices) / len(vertices)
            return (page_index, y_avg)
        except:
            return (0, 0.0)

    # Identifikuojame "Name" entitus (kurie kuria eilutes)
    name_types = ['product_name', 'line_item/description', 'description']
    
    name_entities = [e for e in all_entities if e.type_ in name_types]
    name_entities.sort(key=get_sort_key)
    
    # CLUSTERING NAME ENTITIES BY Y-COORDINATE
    # If multiple name entities are on the same line (e.g. "Barcelo" and "Imperial"), group them into one row.
    rows = []
    if name_entities:
        current_row_entities = [name_entities[0]]
        current_page = get_sort_key(name_entities[0])[0]
        current_y = get_sort_key(name_entities[0])[1]
        
        NAME_MERGE_TOLERANCE = 0.01 # Reduced from 0.02 to 0.01 to avoid merging separate lines
        
        for i in range(1, len(name_entities)):
            e = name_entities[i]
            page, y = get_sort_key(e)
            
            if page == current_page and abs(y - current_y) < NAME_MERGE_TOLERANCE:
                current_row_entities.append(e)
                logging.info(f"Clustering: Merging '{e.mention_text}' into row with '{current_row_entities[0].mention_text}' (Diff: {abs(y - current_y):.4f})")
            else:
                # Finalize current row
                # Calculate average Y for the row
                avg_y = sum(get_sort_key(ent)[1] for ent in current_row_entities) / len(current_row_entities)
                rows.append({
                    'sort_key': (current_page, avg_y),
                    'y': avg_y,
                    'page': current_page,
                    'entities': current_row_entities,
                    'has_name': True
                })
                # Start new row
                current_row_entities = [e]
                current_page = page
                current_y = y
        
        # Append last row
        if current_row_entities:
            avg_y = sum(get_sort_key(ent)[1] for ent in current_row_entities) / len(current_row_entities)
            rows.append({
                'sort_key': (current_page, avg_y),
                'y': avg_y,
                'page': current_page,
                'entities': current_row_entities,
                'has_name': True
            })

    # Prijungiame kitus entitus
    other_entities = [e for e in all_entities if e not in name_entities]
    other_entities.sort(key=get_sort_key)
    
    ASSIGN_TOLERANCE = 0.050 # 5% puslapio aukščio
    
    for e in other_entities:
        e_sort_key = get_sort_key(e)
        e_page = e_sort_key[0]
        e_y = e_sort_key[1]
        
        best_row = None
        min_dist = float('inf')
        
        for row in rows:
            if row['page'] != e_page: continue
            dist = abs(row['y'] - e_y)
            if dist < ASSIGN_TOLERANCE and dist < min_dist:
                min_dist = dist
                best_row = row
        
        if best_row:
            best_row['entities'].append(e)
        else:
            # Našlaitė
            found = False
            for row in rows:
                if row['page'] != e_page: continue
                if not row['has_name']:
                    dist = abs(row['y'] - e_y)
                    if dist < ASSIGN_TOLERANCE:
                        row['entities'].append(e)
                        found = True
                        break
            if not found:
                rows.append({
                    'sort_key': e_sort_key,
                    'y': e_y,
                    'page': e_page,
                    'entities': [e],
                    'has_name': False
                })
    
    rows.sort(key=lambda r: r['sort_key'])
    logging.info(f"Sugrupuota {len(rows)} eilučių.")

    # 4. Konvertuojame į produktų objektus
    products = []
    orphan_rows = []

    for row in rows:
        product = {}
        row_entities = row['entities']
        
        product['_meta_page'] = row.get('page', 0)
        product['_meta_y'] = row.get('y', 0.0)
        
        # Helperiai reikšmių traukimui
        def get_val(types):
            for t in types:
                ent = next((e for e in row_entities if e.type_ == t), None)
                if ent:
                    if hasattr(ent, 'normalized_value') and ent.normalized_value:
                        if hasattr(ent.normalized_value, 'text') and ent.normalized_value.text:
                            return ent.normalized_value.text
                    return ent.mention_text
            return None

        # Name
        # FIX: Removed 'line_item/product_code' and 'product_code' to prevent product numbers from appearing in the name
        name_ents = [e for e in row_entities if e.type_ in ['product_name', 'line_item/description', 'description']]
        # Sort by X coordinate to ensure correct order (e.g. "Barcelo" then "Imperial")
        def get_x(e):
            try:
                page_ref = e.page_anchor.page_refs[0]
                vertices = page_ref.bounding_poly.normalized_vertices
                return vertices[0].x if vertices else 0.0
            except: return 0.0
        name_ents.sort(key=get_x)
        
        name_val = " ".join([e.mention_text for e in name_ents if e.mention_text])
        if name_val:
            product['name'] = name_val.strip()
            
        # Quantity
        q_ent = next((e for e in row_entities if e.type_ in ['quantity', 'line_item/quantity']), None)
        if q_ent:
             val = q_ent.normalized_value.text if hasattr(q_ent, 'normalized_value') and q_ent.normalized_value else q_ent.mention_text
             product['quantity'] = clean_and_convert_to_float(val)

        # Unit Price
        p_ent = next((e for e in row_entities if e.type_ in ['unit_price', 'line_item/unit_price']), None)
        if p_ent:
             val = p_ent.normalized_value.text if hasattr(p_ent, 'normalized_value') and p_ent.normalized_value else p_ent.mention_text
             product['unit_price'] = clean_and_convert_to_float(val)

        # Amount
        a_ent = next((e for e in row_entities if e.type_ in ['amount', 'line_item/amount']), None)
        if a_ent:
             val = a_ent.normalized_value.text if hasattr(a_ent, 'normalized_value') and a_ent.normalized_value else a_ent.mention_text
             product['amount'] = clean_and_convert_to_float(val)
             
        # Volume
        # FIX: Added lowercase 'volume' since Google may return it with different casing
        v_ent = next((e for e in row_entities if e.type_ in ['Volume', 'volume', 'line_item/volume']), None)
        
        # DEBUG: Log all entity types in this row to understand what Google is returning
        row_entity_types = [e.type_ for e in row_entities]
        logging.info(f"DEBUG Row entities for '{product.get('name', 'NO_NAME')[:30]}': {row_entity_types}")
        
        # Check for quantity entity that might actually be volume (unit = liter)
        if not v_ent:
             q_ent_vol = next((e for e in row_entities if e.type_ in ['quantity', 'line_item/quantity'] and 
                               hasattr(e, 'normalized_value') and 
                               hasattr(e.normalized_value, 'text') and 
                               ('l' in e.normalized_value.text.lower() or 'liter' in e.normalized_value.text.lower())), None)
             if q_ent_vol:
                 v_ent = q_ent_vol
                 logging.info(f"Using quantity entity as volume: {v_ent.mention_text}")

        if v_ent:
             val = v_ent.normalized_value.text if hasattr(v_ent, 'normalized_value') and v_ent.normalized_value else v_ent.mention_text
             product['volume'] = clean_volume_value(val)
             logging.info(f"Volume extracted from entity for '{product.get('name', 'NO_NAME')[:30]}': {product['volume']} (raw: '{val}')")
        else:
             logging.info(f"NO Volume entity found for '{product.get('name', 'NO_NAME')[:30]}'! Will use fallback.")

        # ABV
        abv_ent = next((e for e in row_entities if e.type_ in ['abv', 'line_item/abv']), None)
        if abv_ent:
             val = abv_ent.normalized_value.text if hasattr(abv_ent, 'normalized_value') and abv_ent.normalized_value else abv_ent.mention_text
             product['abv'] = clean_and_convert_to_float(val)

        # --- LOGIKA IŠ SENOS FUNKCIJOS (Packaging, Ghost, Volume/ABV recovery) ---
        
        # 1. Packaging check
        try:
            full_line_text = extract_all_text_in_range(document, product.get('_meta_page', 0), product.get('_meta_y', 0), tolerance=0.02)
        except: full_line_text = ""
        
        full_line_lower = full_line_text.lower()
        name_lower = product.get('name', '').lower()
        
        # Discount check
        if any(k in full_line_lower for k in ['discount', 'nuolaida', 'rebate']) or \
           any(k in name_lower for k in ['discount', 'nuolaida', 'rebate']):
             continue

        # Surcharge check (Individual bottle surcharge)
        if any(k in name_lower for k in ['surcharge', 'bottle surcharge', 'individual bottle surcharge']):
             logging.info(f"Eilutė ignoruojama kaip priemoka: {product.get('name')}")
             continue

        # Packaging check
        # FIX: Navimer Alcohol Pur Glass (96%) neturi būti laikomas pakuote, nors turi žodį "Glass"
        is_strong_alcohol = any(k in name_lower for k in ['alcohol pur', 'pure alcohol', '96%', 'spirit', 'navimer'])

        # PATAISYTA: Tikrinti TIK produkto pavadinimą, ne visą eilutę
        # Nes eilutėje gali būti "Bottle" ar kitų žodžių iš kitu produktų
        name_has_packaging = any(k in name_lower for k in PACKAGING_KEYWORDS)
        
        # Patikrinti ar produktas atrodo kaip vynas/alkoholis pagal pavadinimą
        # Jei taip - neturėtų būti laikomas pakuote net jei ABV=0
        wine_indicators = [
            'wine', 'vino', 'vin ', 'wein', 'rouge', 'blanc', 'rosso', 'bianco', 'rose', 'rosé',
            'cabernet', 'merlot', 'chardonnay', 'sauvignon', 'pinot', 'shiraz', 'syrah', 'riesling',
            'malbec', 'tempranillo', 'sangiovese', 'nebbiolo', 'barbera', 'primitivo', 'zinfandel',
            'grenache', 'mourvedre', 'viognier', 'gewurztraminer', 'gruner', 'verdejo', 'albarino',
            'vermentino', 'trebbiano', 'garganega', 'corvina', 'montepulciano', 'nero d\'avola',
            'champagne', 'prosecco', 'cava', 'cremant', 'sekt', 'spumante', 'franciacorta',
            'rioja', 'chianti', 'barolo', 'barbaresco', 'brunello', 'amarone', 'valpolicella',
            'bordeaux', 'burgundy', 'bourgogne', 'chablis', 'sancerre', 'pouilly', 'cotes du rhone',
            'chateau', 'domaine', 'estate', 'reserve', 'reserva', 'gran reserva', 'crianza',
            'doc', 'docg', 'igt', 'igp', 'aoc', 'aop', 'dop',
            # Vyno regionai
            'toscana', 'tuscany', 'piemonte', 'veneto', 'sicilia', 'puglia', 'lombardia',
            'champagne', 'alsace', 'loire', 'provence', 'languedoc', 'rhone',
            'napa', 'sonoma', 'mendoza', 'marlborough', 'barossa', 'stellenbosch',
            # Vyno gamintojai
            'astruc', 'antinori', 'frescobaldi', 'gaja', 'sassicaia', 'tignanello',
            # Putojantys
            'brut', 'extra brut', 'demi sec', 'millesime'
        ]
        looks_like_wine = any(k in name_lower for k in wine_indicators)
        
        # PATAISYTA: Produktas yra pakuotė TIK jei pavadinime yra AIŠKUS packaging raktažodis
        # Nebedaryti prielaidų kad produktas be tūrio/ABV yra pakuotė - tai blokuoja akcizų skaičiavimą
        # Produktas NĖRA pakuotė jei:
        # - Atrodo kaip vynas/alkoholis (wine_indicators)
        # - Yra stiprus alkoholis (is_strong_alcohol)
        is_packaging = name_has_packaging and not is_strong_alcohol and not looks_like_wine
        
        if is_packaging:
            product['volume'] = 0.0
            product['abv'] = 0.0
            product['_is_packaging'] = True  # Žymime kaip pakuotę, kad ABV recovery neveiktų
            logging.info(f"Pakuotė: {product.get('name')}")

        # Našlaitės - eilutės BE pavadinimo, bet SU skaičiais
        if not product.get('name') and any(product.get(k) for k in ['quantity', 'unit_price', 'amount']):
            orphan_rows.append(product)
            continue
        
        # NAUJA: "Pavadinimo tęsiniai" - eilutės SU pavadinimu, bet BE skaičių
        # Pvz. "Der Rheinberger Kräuterbitter" po "Underberg 4/30" eilutės
        has_numbers = any(product.get(k) and product.get(k) > 0 for k in ['quantity', 'unit_price', 'amount'])
        if product.get('name') and not has_numbers:
            orphan_rows.append(product)
            product['_is_name_continuation'] = True  # Žymime kaip pavadinimo tęsinį
            logging.info(f"Pavadinimo tęsinys (be skaičių): '{product.get('name')}'")
            continue

        if product.get('name'):
            # Ghost check - PERKELTAS Į PABAIGĄ
            # Anksčiau čia tikrindavome ir trindavome eilutes be skaičių.
            # Dabar leidžiame joms likti, kad "našlaičių poravimas" galėtų užpildyti trūkstamus duomenis.
            
            # Volume recovery - TIK specialūs butelių dydžiai pagal pavadinimą
            # SVARBU: Nenaudojame "surrounding text" nes tai gali paimti neteisingą tūrį iš gretimos eilutės
            if not product.get('volume'):
                name_lower = product['name'].lower()
                product_name = product['name']
                
                # NAUJA: Pirma bandome ištraukti tūrį iš pavadinimo (pvz. "0,02 Underberg")
                extracted_vol = extract_volume_from_text(product_name)
                if extracted_vol:
                    product['volume'] = extracted_vol
                    logging.info(f"Tūris ištrauktas iš pavadinimo '{product_name[:40]}': {extracted_vol}L")
                elif any(k in name_lower for k in ['packaging', 'pakuotė', 'dėžutė', 'empty box', 'empty gift box']):
                    product['volume'] = 0.0
                elif 'double magnum' in name_lower: product['volume'] = 3.0
                elif 'magnum' in name_lower: product['volume'] = 1.5
                elif 'jeroboam' in name_lower: product['volume'] = 3.0
                elif 'rehoboam' in name_lower: product['volume'] = 4.5
                elif 'mathusalem' in name_lower: product['volume'] = 6.0
                elif 'salmanazar' in name_lower: product['volume'] = 9.0
                elif 'balthazar' in name_lower: product['volume'] = 12.0
                elif 'nebuchadnezzar' in name_lower: product['volume'] = 15.0
                elif '100cl' in name_lower: product['volume'] = 1.0
                else:
                    # Jei tūris nerastas, naudojame standartinį 0.75L vynui/alkoholiui
                    # Tai leidžia akcizo skaičiavimui veikti, vartotojas gali pataisyti jei neteisinga
                    product['volume'] = 0.75
                    logging.warning(f"Tūris nerastas produktui '{product['name'][:40]}' - naudojamas standartinis 0.75L")

            # ABV recovery - TIK jei tai NE pakuotė
            # Pakuotėms (gift box ir pan.) neieškome ABV nes Google grąžino 0%
            if not product.get('abv') and not product.get('_is_packaging'):
                rec_abv = recover_abv_from_surrounding_text(document, product.get('_meta_page', 0), product.get('_meta_y', 0))
                if rec_abv: product['abv'] = rec_abv
                else: product['abv'] = estimate_abv_from_name(product['name'])
            elif product.get('_is_packaging'):
                product['abv'] = 0.0  # Pakuotėms visada 0%

            products.append(product)

    # Našlaičių poravimas
    if orphan_rows:
        logging.info(f"Pradedamas našlaičių ({len(orphan_rows)}) poravimas.")
        for orphan in orphan_rows:
            best_target = None
            min_dist = float('inf')
            orphan_page = orphan.get('_meta_page', 0)
            orphan_y = orphan.get('_meta_y', 0.0)
            is_name_continuation = orphan.get('_is_name_continuation', False)
            
            for prod in products:
                # needs_data check removed to allow merging text/details into full products
                # But we must prevent merging two full products (see below)
                if prod.get('_meta_page', 0) != orphan_page: continue
                
                dist = orphan_y - prod.get('_meta_y', 0.0)
                # Pavadinimo tęsiniams: ieškome produkto VIRŠ (dist > 0 = našlaitė yra žemiau)
                # Padidintas diapazonas iki 0.15 nes pavadinimas gali būti per kelias eilutes
                if is_name_continuation:
                    if 0.0 < dist <= 0.15:  # Našlaitė yra ŽEMIAU produkto
                        if abs(dist) < min_dist:
                            min_dist = abs(dist)
                            best_target = prod
                else:
                    # Standartinė logika kitoms našlaitėms
                    if -0.01 <= dist <= 0.10:
                        if abs(dist) < min_dist:
                            min_dist = abs(dist)
                            best_target = prod
            
            # Check if both are "Full Products" (have Quantity AND (Price OR Amount))
            orphan_is_full = (
                (orphan.get('quantity') or 0) > 0 and 
                ((orphan.get('unit_price') or 0) > 0 or (orphan.get('amount') or 0) > 0)
            )
            
            if best_target:
                target_is_full = (
                    (best_target.get('quantity') or 0) > 0 and 
                    ((best_target.get('unit_price') or 0) > 0 or (best_target.get('amount') or 0) > 0)
                )
                
                # If both are full products, DO NOT MERGE. Treat orphan as separate product.
                if orphan_is_full and target_is_full:
                    logging.info(f"Našlaitė turi pilnus duomenis (Q={orphan.get('quantity')}) ir taikinys taip pat. Neporuojama, laikoma atskiru produktu.")
                    best_target = None

            if best_target:
                for key in ['quantity', 'unit_price', 'amount']:
                    if best_target.get(key) in [None, 0, 0.0] and orphan.get(key) is not None:
                        best_target[key] = orphan[key]
                
                # PATAISYTA: Jei tai pavadinimo tęsinys, naudojame orphan['name'] tiesiogiai
                if is_name_continuation and orphan.get('name'):
                    current_name = best_target.get('name', '')
                    orphan_name = orphan.get('name', '')
                    if orphan_name and orphan_name not in current_name:
                        best_target['name'] = (current_name + " " + orphan_name).strip()
                        logging.info(f"Prijungtas pavadinimo tęsinys prie '{current_name}': '{orphan_name}'")
                else:
                    # Prijungiame tekstą iš našlaitės eilutės prie produkto pavadinimo
                    # Tai padeda atvejais, kai pavadinimas yra per kelias eilutes (pvz. Oloroso)
                    orphan_text = extract_text_at_row(document, orphan_page, orphan_y, tolerance=0.02)
                    if orphan_text:
                        current_name = best_target.get('name', '')
                        # Tikriname ar tekstas jau nėra pavadinime (kad nedubliuotume)
                        # Taip pat ignoruojame jei tekstas yra tik skaičiai
                        if orphan_text not in current_name and not re.match(r'^[\d\s.,€]+$', orphan_text):
                            best_target['name'] = (current_name + " " + orphan_text).strip()
                            logging.info(f"Prijungtas tekstas iš našlaitės prie '{current_name}': '{orphan_text}'")

                    # --- NAUJA: Bandome atnaujinti ABV ir Tūrį iš našlaitės teksto ---
                    # Dažnai ABV būna antroje eilutėje (pvz. Oloroso ... 46%)
                    orphan_abv = extract_abv_from_text(orphan_text)
                    if orphan_abv:
                        current_abv = best_target.get('abv', 0.0)
                        # Atnaujiname jei ABV nerastas arba jei našlaitės ABV atrodo patikimesnis (pvz. ne 0)
                        # Arba jei esamas ABV yra "standartinis" (pvz. 40.0), o našlaitėje specifinis (46.0)
                        if current_abv == 0.0 or (orphan_abv != current_abv and orphan_abv > 0):
                            best_target['abv'] = orphan_abv
                            logging.info(f"Atnaujintas ABV iš našlaitės teksto: {orphan_abv}% (buvo {current_abv}%)")

                    orphan_vol = extract_volume_from_text(orphan_text)
                    if orphan_vol:
                        current_vol = best_target.get('volume', 0.0)
                        if current_vol == 0.0:
                            best_target['volume'] = orphan_vol
                            logging.info(f"Atnaujintas tūris iš našlaitės teksto: {orphan_vol}")
                    # ----------------------------------------------------------------

                # Taip pat patikriname ar našlaitė turi entity-based ABV/Volume
                if orphan.get('abv') and (best_target.get('abv') in [None, 0, 0.0]):
                    best_target['abv'] = orphan['abv']
                if orphan.get('volume') and (best_target.get('volume') in [None, 0, 0.0]):
                    best_target['volume'] = orphan['volume']

                # Recalc unit price
                qty = best_target.get('quantity') or 0
                amt = best_target.get('amount') or 0
                if qty and amt and (best_target.get('unit_price') in [None, 0, 0.0]):
                    best_target['unit_price'] = amt / qty
            else:
                # New product from orphan
                if orphan.get('unit_price') in [None, 0, 0.0] and orphan.get('amount') in [None, 0, 0.0]: continue
                if is_pallet_or_packaging_row(document, orphan_page, orphan_y): continue
                
                if not orphan.get('name'):
                    rec_name = extract_text_at_row(document, orphan_page, orphan_y, tolerance=0.05)
                    orphan['name'] = rec_name if rec_name else '(be pavadinimo)'
                
                # Surcharge check for orphans
                if any(k in orphan['name'].lower() for k in ['surcharge', 'bottle surcharge', 'individual bottle surcharge']):
                     logging.info(f"Našlaitė ignoruojama kaip priemoka: {orphan.get('name')}")
                     continue

                # Ensure defaults
                if not orphan.get('volume'): orphan['volume'] = 0.0
                if not orphan.get('abv'): orphan['abv'] = 0.0
                
                products.append(orphan)

    # Final cleanup of empty products (Ghosts)
    # Dabar, kai našlaitės jau prijungtos, galime saugiai išvalyti produktus, kurie vis dar neturi duomenų.
    final_products = []
    if products:
        final_products.append(products[0])
        for i in range(1, len(products)):
            current = products[i]
            prev = final_products[-1]
            
            is_empty = (
                (current.get('quantity') in [None, 0, 0.0]) and 
                (current.get('unit_price') in [None, 0, 0.0]) and 
                (current.get('amount') in [None, 0, 0.0])
            )
            
            if is_empty:
                # Check distance to previous
                if current.get('_meta_page') == prev.get('_meta_page'):
                    dist = current.get('_meta_y', 0.0) - prev.get('_meta_y', 0.0)
                    if 0 < dist < 0.06:
                        logging.info(f"Sujungiamas 'Ghost' produktas: '{current.get('name')}' prijungiamas prie '{prev.get('name')}'")
                        prev['name'] += " " + current.get('name', '')
                        continue # Skip adding current to final_products
            
            final_products.append(current)
    
    products = final_products

    # Final amount check
    for prod in products:
        if prod.get('amount') in [None, 0, 0.0]:
            qty = prod.get('quantity') or 0
            price = prod.get('unit_price') or 0
            if qty > 0 and price > 0:
                prod['amount'] = qty * price

    return products

def extract_line_items_from_doc_ai_legacy(document) -> list:
    """Ištraukia prekių eilutes iš Document AI response (senas, nepatikimas metodas)."""
    products = []
    for page in document.pages:
        for table in page.tables:
            for row in table.body_rows:
                product = {}
                row_text = []
                for cell in row.cells:
                    cell_text = "".join(
                        [
                            document.text[segment.start_index : segment.end_index]
                            for segment in cell.layout.text_anchor.text_segments
                        ]
                    )
                    row_text.append(cell_text)

                # This is still a very brittle way to do this, but it's the only way without more info
                # on the structure of the tables.
                # We assume that the product name is the first long string, and the rest are numbers.
                for text in row_text:
                    # Patobulintas tikrinimas: ieškome teksto, kuris yra ilgesnis nei 8 simboliai ir TURI bent vieną raidę.
                    if not product.get("name") and len(text) > 8 and re.search(r'[a-zA-Z]', text):
                        product["name"] = text
                    elif not product.get("quantity") and text.isdigit():
                        product["quantity"] = clean_and_convert_to_float(text)
                    elif not product.get("unit_price") and re.match(r"^\d+,\d+$", text):
                        product["unit_price"] = clean_and_convert_to_float(text)
                    elif not product.get("amount") and re.match(r"^\d+\.\d+,\d+$", text):
                        product["amount"] = clean_and_convert_to_float(text)

                # Handle volume
                if product.get("name") and "magnum" in product["name"].lower():
                    product["volume"] = 1.5
                elif not product.get("volume_l"):
                    product["volume_l"] = 0.75

                if product:
                    products.append(product)
    return products

def extract_summary_with_deepseek(document_text: str) -> dict:
    """Specializuota funkcija, ieškanti nuolaidos ir transporto."""
    logging.info("AI_INVOICE.PY - (extract_summary) Ieškoma nuolaidos ir transporto su DeepSeek...")
    if not deepseek_client: return {"discount_amount": 0.0, "transport_amount": 0.0}
    
    prompt_content = f"""Išanalizuok šį sąskaitos tekstą ir surask:
1. Bendrą nuolaidos sumą - ieškok eilučių su raktažodžiais: 'ESCOMPTE', 'Discount', 'Nuolaida', 'Rabatt', 'Remise', 'Descuento', 'Sconto'
2. Transporto išlaidas - ieškok eilučių su raktažodžiais: 'Freight', 'Transport', 'Fracht', 'Livraison', 'Spedizione', 'Vracht', 'Transporte', 'Shipping', 'Delivery'

SVARBU: IGNORUOK šiuos produktus, net jei jie turi transporto raktažodžius:
- 'Pallets' arba 'Pallet' - tai yra produktas, ne transporto išlaidos
- Bet kokie alkoholiniai gėrimai (vynas, alus, whiskey, vodka ir pan.)

Grąžink JSON objektą su dviem raktais: 'discount_amount' ir 'transport_amount'. 
Reikšmės turi būti skaičiai (nuolaida gali būti neigiama). 
Jei nerandi, grąžink 0.0.

Pavyzdys: {{"discount_amount": -50.0, "transport_amount": 150.0}}

Tekstas:
---
{document_text[:30000]}
---"""
    
    payload = {
        "model": "deepseek-chat", 
        "messages": [
            {"role": "system", "content": "Tu esi finansų analitikas. Tavo užduotis - rasti nuolaidą ir transporto išlaidas sąskaitoje ir grąžinti jas JSON formatu."},
            {"role": "user", "content": prompt_content}
        ], 
        "max_tokens": 300, 
        "temperature": 0.0
    }
    
    try:
        # Pridėtas timeout=10s
        response = deepseek_client.chat.completions.create(**payload, timeout=10)
        content = response.choices[0].message.content
        
        if not content:
            logging.warning("DeepSeek grąžino tuščią atsakymą")
            return {"discount_amount": 0.0, "transport_amount": 0.0}
        
        # Bandome ištraukti JSON
        match = re.search(r"```json\s*([\s\S]*?)\s*```", content)
        json_to_parse = match.group(1).strip() if match else content.strip()
        
        try:
            data = json.loads(json_to_parse)
        except json.JSONDecodeError as json_error:
            logging.warning(f"DeepSeek JSON dekodavimo klaida: {json_error}. Turinys: {content[:200]}...")
            return {"discount_amount": 0.0, "transport_amount": 0.0}
        
        # Validuojame and nustatome default reikšmes
        discount_amount = clean_and_convert_to_float(data.get('discount_amount', 0.0)) or 0.0
        transport_amount = clean_and_convert_to_float(data.get('transport_amount', 0.0)) or 0.0
        
        # Validuojame transporto sumą
        if transport_amount > 0:
            is_valid, error_msg = validate_transport_amount(transport_amount)
            if not is_valid:
                logging.warning(f"DeepSeek transporto validacijos klaida: {error_msg}")
                transport_amount = 0.0
        
        result = {
            "discount_amount": discount_amount,
            "transport_amount": transport_amount
        }
        
        if discount_amount != 0.0:
            logging.info(f"DeepSeek rado nuolaidą: {discount_amount}")
        if transport_amount > 0.0:
            logging.info(f"DeepSeek rado transportą: {transport_amount}")
            
        return result
        
    except openai.APIConnectionError as e:
        logging.error(f"DeepSeek API ryšio klaida: {e}")
        return {"discount_amount": 0.0, "transport_amount": 0.0}
    except openai.RateLimitError as e:
        logging.error(f"DeepSeek API limito klaida: {e}")
        return {"discount_amount": 0.0, "transport_amount": 0.0}
    except openai.APIError as e:
        logging.error(f"DeepSeek API klaida: {e}")
        return {"discount_amount": 0.0, "transport_amount": 0.0}
    except Exception as e:
        logging.error(f"AI_INVOICE.PY - Netikėta klaida (extract_summary): {e}", exc_info=True)
        return {"discount_amount": 0.0, "transport_amount": 0.0}

def check_document_ai_availability() -> tuple[bool, str]:
    """
    Patikrina ar Document AI yra pasiekiamas.
    
    Returns:
        tuple: (ar_pasiekiamas, klaidos_pranešimas)
    """
    try:
        credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        if not credentials_path:
            return False, "GOOGLE_APPLICATION_CREDENTIALS kintamasis nenustatytas"
        
        if not os.path.exists(credentials_path):
            return False, f"Prisijungimo duomenų failas nerastas: {credentials_path}"
        
        # Bandome sukurti credentials
        credentials = service_account.Credentials.from_service_account_file(credentials_path)
        if not credentials:
            return False, "Nepavyko sukurti Google Cloud credentials"
        
        # Patikrinti ar nustatyti reikalingi kintamieji
        project_id = os.getenv("DOCAI_PROJECT_ID")
        processor_id = os.getenv("DOCAI_PROCESSOR_ID")
        
        if not project_id:
            return False, "DOCAI_PROJECT_ID kintamasis nenustatytas"
        if not processor_id:
            return False, "DOCAI_PROCESSOR_ID kintamasis nenustatytas"
        
        return True, "Document AI konfigūracija tinkama"
        
    except Exception as e:
        return False, f"Document AI konfigūracijos klaida: {str(e)}"

# --- PAGRINDINĖ FUNKCIJA ---
def extract_invoice_data(file_path: str, manual_transport: float = 0.0) -> dict:
    try:
        log_function_call("extract_invoice_data", file_path=file_path, manual_transport=manual_transport)
        
        # 1. Tikrinti cache
        cache_key = f"{file_path}_{manual_transport}"
        cached_result = pdf_cache.get(cache_key)
        if cached_result:
            logging.info("Naudojami cache duomenys: %s", file_path)
            return cached_result
        
        # 2. Patikrinti Document AI prieinamumą
        is_available, availability_msg = check_document_ai_availability()
        if not is_available:
            logging.error(f"Document AI nepasiekiamas: {availability_msg}")
            return {'error': f'Document AI konfigūracijos klaida: {availability_msg}', 
                   'products': [], 'summary': {'discount_amount': 0.0, 'transport_amount': 0.0}}
        
        logging.info("Document AI konfigūracija patikrinta - viskas tvarkoj")

        # --- VAIZDO APDOROJIMAS (PREPROCESSING) ---
        # Bandome pagerinti vaizdo kokybę prieš siunčiant į Document AI
        processed_file_path = file_path + "_processed.pdf"
        use_processed_file = False
        
        # VARTOTOJO PRAŠYMU IŠJUNGTA
        logging.info("Vaizdo apdorojimas (preprocessing) yra IŠJUNGTAS vartotojo prašymu.")
        # try:
        #     if preprocess_pdf(file_path, processed_file_path):
        #         logging.info(f"Sėkmingai atliktas vaizdo apdorojimas. Naudojamas failas: {processed_file_path}")
        #         use_processed_file = True
        #     else:
        #         logging.info("Vaizdo apdorojimas nepavyko arba buvo praleistas. Naudojamas originalus failas.")
        # except Exception as e:
        #     logging.warning(f"Klaida bandant apdoroti vaizdą: {e}")
        
        file_to_process = processed_file_path if use_processed_file else file_path
        # ------------------------------------------
        
        # --- WORD DOKUMENTŲ KONVERTAVIMAS ---
        converted_pdf = None
        if is_word_document(file_to_process):
            logging.info(f"Aptiktas Word dokumentas: {file_to_process}")
            converted_pdf = convert_word_to_pdf(file_to_process)
            if converted_pdf and os.path.exists(converted_pdf):
                file_to_process = converted_pdf
                logging.info(f"Naudojamas konvertuotas PDF: {file_to_process}")
            else:
                logging.error("Nepavyko konvertuoti Word dokumento į PDF")
                return {'error': 'Nepavyko konvertuoti Word dokumento. Bandykite įkelti PDF arba paveikslėlį.', 
                       'products': [], 'summary': {'discount_amount': 0.0, 'transport_amount': 0.0}}
        
        # Nustatome MIME tipą pagal failo plėtinį
        mime_type = get_mime_type(file_to_process)
        logging.info(f"Failo MIME tipas: {mime_type}")
        # ------------------------------------------
        
        # 2. Prisijungimas ir kreipimasis į Document AI
        credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        credentials = service_account.Credentials.from_service_account_file(credentials_path)
        
        doc_ai_location = os.getenv("DOCAI_LOCATION", "us")
        doc_ai_client = documentai.DocumentProcessorServiceClient(credentials=credentials, client_options={"api_endpoint": f"{doc_ai_location}-documentai.googleapis.com"})
        
        processor_path = doc_ai_client.processor_path(os.getenv("DOCAI_PROJECT_ID"), doc_ai_location, os.getenv("DOCAI_PROCESSOR_ID"))

        with open(file_to_process, "rb") as input_file:
            raw_document = documentai.RawDocument(content=input_file.read(), mime_type=mime_type)
        
        result = doc_ai_client.process_document(request=documentai.ProcessRequest(name=processor_path, raw_document=raw_document))
        document = result.document

        # Išvalome laikiną failą
        if use_processed_file and os.path.exists(processed_file_path):
            try:
                os.remove(processed_file_path)
            except:
                pass
        
        # Išvalome konvertuotą PDF failą
        if converted_pdf and os.path.exists(converted_pdf):
            try:
                os.remove(converted_pdf)
                logging.info(f"Ištrintas konvertuotas PDF: {converted_pdf}")
            except:
                pass

        logging.info("--- Document AI Full Response ---")
        logging.info(f"Document text: {document.text}")
        for entity in document.entities:
            logging.info(f"Entity type: {entity.type_}, mention_text: {entity.mention_text}, confidence: {entity.confidence}")
        logging.info("--- End of Document AI Full Response ---")

        if not document.text:
            logging.warning("Document AI negrąžino dokumento teksto.")
            return {}
        
        # 2. TRANSPORTO APTIKIMAS IŠ DOCUMENT AI
        document_ai_transport = safe_extract_transport(document)
        
        # 3. PRODUKTŲ IŠTRAUKIMAS (naudojam patikimą metodą)
        products = extract_line_items_from_doc_ai(document)
        logging.info(f"Ištraukta produktų eilučių: {len(products) if products else 0}")
        
        # 4. NUOLAIDOS IR TRANSPORTO IŠTRAUKIMAS SU DEEPSEEK
        summary = extract_summary_with_deepseek(document.text)
        logging.info(f"Summary rezultatas: {summary}")
        
        # 5. TRANSPORTO EILUČIŲ FILTRAVIMAS IŠ PRODUKTŲ
        if products:
            filtered_products, products_transport = filter_transport_lines(products)
            logging.info(f"Po filtravimo produktų: {len(filtered_products)}, transportas iš produktų: {products_transport}")
            products = filtered_products
        else:
            products_transport = 0.0
            logging.warning("Produktų sąrašas tuščias - filtravimas praleistas")

        if not products:
             logging.error("Nepavyko ištraukti prekių eilučių.")
             # Vis tiek grąžiname transporto informaciją su teisingais prioritetais
             deepseek_transport = summary.get('transport_amount', 0.0)
             
             # Prioritetų logika net ir be produktų
             if manual_transport > 0:
                 total_transport = manual_transport
                 transport_source = "manual"
                 logging.info(f"Naudojamas rankinis transportas: {total_transport} EUR")
             else:
                 auto_transport = max(document_ai_transport, deepseek_transport, products_transport)
                 if auto_transport > 0:
                     total_transport = auto_transport
                     transport_source = "automatic"
                     logging.info(f"Naudojamas automatiškai rastas transportas: {total_transport} EUR (Document AI: {document_ai_transport}, DeepSeek: {deepseek_transport}, Produktai: {products_transport})")
                 else:
                     total_transport = 0.0
                     transport_source = "none"
                     logging.info("Transportas nerastas")
             
             summary['transport_amount'] = total_transport
             summary['transport_source'] = transport_source
             return {'products': [], 'summary': summary}

        # 6. BENDROS TRANSPORTO SUMOS NUSTATYMAS SU PRIORITETŲ LOGIKA
        deepseek_transport = summary.get('transport_amount', 0.0)
        
        # Ištraukiame tiekėjo pavadinimą iš Document AI response
        supplier_name = 'Unknown Supplier'
        try:
            # Document AI entities sąraše ieškome tiekėjo pavadinimo
            for entity in document.entities:
                if entity.type_ == 'supplier_name':
                    supplier_name = entity.mention_text.strip()
                    logging.info(f"Document AI rado tiekėjo pavadinimą: {supplier_name}")
                    break
        except Exception as e:
            logging.warning(f"Nepavyko ištraukti tiekėjo pavadinimo iš Document AI: {e}")
        
        summary['supplier_name'] = supplier_name

        # Prioritetų logika:
        # 1. Rankinis įvedimas (jei > 0)
        # 2. Automatinis aptikimas (didžiausia suma iš visų šaltinių)
        # 3. Nulis
        if manual_transport > 0:
            total_transport = manual_transport
            transport_source = "manual"
            logging.info(f"Naudojamas rankinis transportas: {total_transport} EUR")
        else:
            auto_transport = max(document_ai_transport, deepseek_transport, products_transport)
            if auto_transport > 0:
                total_transport = auto_transport
                transport_source = "automatic"
                logging.info(f"Naudojamas automatiškai rastas transportas: {total_transport} EUR (Document AI: {document_ai_transport}, DeepSeek: {deepseek_transport}, Produktai: {products_transport})")
            else:
                total_transport = 0.0
                transport_source = "none"
                logging.info("Transportas nerastas")
        
        # Atnaujinome summary su galutine transporto suma ir šaltiniu
        summary['transport_amount'] = total_transport
        summary['transport_source'] = transport_source

        # 7. Produktų apdorojimas ir klasifikavimas
        processed_products = []
        for item_data in products:
            product_name = item_data.get('name', '')
            
            # --- PAVADINIMO VALYMAS ---
            # Naudojame clean_product_name funkciją, kuri pašalina pakavimo informaciją
            product_name = clean_product_name(product_name)
            
            extracted_abv = clean_and_convert_to_float(item_data.get('abv'))
            extracted_volume = clean_and_convert_to_float(item_data.get('volume'))
            
            # Patikrinti ar tai pakuotė (gift box ir pan.) - jei volume=0 IR abv=0, tai pakuotė
            # BET tik jei produktas NEATRODO kaip vynas/alkoholis
            is_packaging = item_data.get('_is_packaging', False)
            if not is_packaging and extracted_volume == 0 and extracted_abv == 0:
                # Papildomas patikrinimas pagal pavadinimą
                name_lower = product_name.lower()
                
                # Patikrinti ar produktas atrodo kaip vynas
                wine_indicators = [
                    'wine', 'vino', 'vin ', 'wein', 'rouge', 'blanc', 'rosso', 'bianco', 'rose', 'rosé',
                    'cabernet', 'merlot', 'chardonnay', 'sauvignon', 'pinot', 'shiraz', 'syrah', 'riesling',
                    'vermentino', 'trebbiano', 'garganega', 'montepulciano', 'nero d\'avola',
                    'champagne', 'prosecco', 'cava', 'brut', 'rioja', 'chianti', 'barolo', 'doc', 'docg', 'igt', 'igp',
                    'astruc', 'antinori', 'frescobaldi'
                ]
                looks_like_wine = any(k in name_lower for k in wine_indicators)
                
                # Tik jei atrodo kaip pakuotė IR NEATRODO kaip vynas
                if any(k in name_lower for k in ['gift box', 'giftbox', 'empty box', 'packaging', 'pakuotė', 'wooden box']):
                    is_packaging = True
                elif not looks_like_wine:
                    # Jei neatrodo kaip vynas ir volume=0, abv=0 - gali būti pakuotė
                    # Bet jei atrodo kaip vynas - tai vynas be ABV nurodymo
                    is_packaging = False  # Paliekame kaip False, nes gali būti vynas be ABV
                    is_packaging = True
            
            # Jei ABV nerastas arba yra 0, bandome nustatyti pagal produkto pavadinimą
            # BET TIK jei tai NE pakuotė!
            if not extracted_abv or extracted_abv == 0.0:
                if is_packaging:
                    # Pakuotėms visada 0% - negalime "atspėti" ABV
                    final_abv = 0.0
                    logging.info(f"Pakuotė '{product_name}' - ABV nustatytas 0%")
                else:
                    estimated_abv = estimate_abv_from_name(product_name)
                    if estimated_abv > 0:
                        logging.info(f"ABV nerastas produktui '{product_name}', nustačiau tipinį: {estimated_abv}%")
                        final_abv = estimated_abv
                    else:
                        final_abv = 0.0
            else:
                final_abv = extracted_abv
                logging.info(f"ABV rastas produktui '{product_name}': {final_abv}%")
            
            item_copy = {
                'name': product_name,
                'quantity': clean_and_convert_to_float(item_data.get('quantity')),
                'unit_price': clean_and_convert_to_float(item_data.get('unit_price')),
                'amount': clean_and_convert_to_float(item_data.get('amount')),
                'volume': clean_and_convert_to_float(item_data.get('volume')),
                'abv': final_abv
            }
            
            # DEBUG: Log volume info
            logging.info(f"DEBUG VOLUME: '{product_name[:40]}' - raw volume from Google: {item_data.get('volume')}, converted: {item_copy['volume']}")
            
            if item_copy.get('name'):
                logging.info(f"Klasifikuojamas produktas: '{item_copy['name']}' su ABV: {final_abv}%")
                item_copy['excise_category_key'] = classify_product_with_deepseek(item_copy['name'], final_abv)
                processed_products.append(item_copy)

        result = {'products': processed_products, 'summary': summary}
        
        # Išsaugoti į cache
        pdf_cache.set(cache_key, result)
        
        return result

    except FileNotFoundError as e:
        logging.error(f"PDF failas nerastas: {e}")
        return {'error': 'PDF failas nerastas', 'products': [], 'summary': {'discount_amount': 0.0, 'transport_amount': 0.0}}
    except PermissionError as e:
        logging.error(f"Nėra teisių skaityti PDF failą: {e}")
        return {'error': 'Nėra teisių skaityti failą', 'products': [], 'summary': {'discount_amount': 0.0, 'transport_amount': 0.0}}
    except Exception as e:
        logging.error(f"AI_INVOICE.PY - Netikėta klaida `extract_invoice_data`: {e}", exc_info=True)
        return {'error': f'Sistemos klaida: {str(e)}', 'products': [], 'summary': {'discount_amount': 0.0, 'transport_amount': 0.0}}

def extract_all_text_in_range(document, page_index, y_center, tolerance=0.05) -> str:
    """Ištraukia visą tekstą iš nurodyto vertikalaus diapazono, sujungdamas visas eilutes."""
    try:
        if page_index >= len(document.pages):
            return ""
            
        page = document.pages[page_index]
        items = page.lines if page.lines else page.tokens
        
        lines_text = []
        
        for item in items:
            try:
                vertices = item.layout.bounding_poly.normalized_vertices
                y_item = sum(v.y for v in vertices) / len(vertices)
                
                if abs(y_item - y_center) < tolerance:
                    text = ""
                    for segment in item.layout.text_anchor.text_segments:
                        start = int(segment.start_index)
                        end = int(segment.end_index)
                        text += document.text[start:end]
                    
                    if text.strip():
                        lines_text.append((y_item, text.strip()))
            except:
                continue
        
        # Rūšiuojame pagal Y (iš viršaus į apačią)
        lines_text.sort(key=lambda x: x[0])
        
        return " ".join([t[1] for t in lines_text])
        
    except Exception as e:
        logging.warning(f"Klaida traukiant tekstą diapazone: {e}")
        return ""

def extract_volume_from_text(text: str) -> Optional[float]:
    """Bando ištraukti tūrį iš teksto eilutės (su vienetais arba be)."""
    if not text:
        return None
    
    # 0. NAUJA: Tūris pavadinimo pradžioje (pvz. "0,02 Underberg" arba "0.75 Chianti")
    # Vokiški/prancūziški formatai naudoja kablelį
    match = re.match(r'^(\d+[.,]\d+)\s+[A-Za-z]', text)
    if match:
        try:
            val_str = match.group(1).replace(',', '.')
            val = float(val_str)
            # Tūris turi būti logiškas (0.01 - 20 litrų)
            if 0.01 <= val <= 20.0:
                logging.info(f"Rastas tūris pavadinimo pradžioje: {val}L")
                return val
        except ValueError:
            pass
    
    # 1. Liter / Liters / L / ltr
    # Ieškome skaičiaus prieš vienetą. Pvz: "1 liter", "0.7 l", "1,0 ltr"
    match = re.search(r'\b(\d+(?:[.,]\d+)?)\s*(?:liter|liters|l|cl|ml|ltr)\b', text, re.IGNORECASE)
    if match:
        try:
            val_str = match.group(1).replace(',', '.')
            val = float(val_str)
            # Saugiklis: tūris turi būti logiškas (0.01 - 20 litrų)
            if 0.01 <= val <= 20.0:
                return val
        except ValueError:
            pass

    # 2. cl (centilitrai) -> verčiame į litrus
    match = re.search(r'\b(\d+(?:[.,]\d+)?)\s*cl\b', text, re.IGNORECASE)
    if match:
        try:
            val_str = match.group(1).replace(',', '.')
            val = float(val_str) / 100.0
            if 0.01 <= val <= 20.0:
                return val
        except ValueError:
            pass

    # 3. ml (mililitrai) -> verčiame į litrus
    match = re.search(r'\b(\d+(?:[.,]\d+)?)\s*ml\b', text, re.IGNORECASE)
    if match:
        try:
            val_str = match.group(1).replace(',', '.')
            val = float(val_str) / 1000.0
            if 0.01 <= val <= 20.0:
                return val
        except ValueError:
            pass
            
    # 4. "Naked" skaičiai (standartiniai tūriai be vienetų)
    # Ieškome skaičių, kurie atitinka standartinius tūrius: 0.2, 0.5, 0.7, 0.75, 1.0
    # Svarbu: turi būti atskirti tarpais ar ribomis
    # Naudojame lookbehind/lookahead, kad nepagautume skaičiaus viduryje kito skaičiaus
    matches = re.findall(r'(?<!\d)(\d+(?:[.,]\d+)?)(?!\d)', text)
    for m in matches:
        try:
            val = float(m.replace(',', '.'))
            # Standartiniai butelių dydžiai (L)
            if val in [0.2, 0.5, 0.7, 0.75, 1.0, 1.5, 3.0]:
                 logging.info(f"Rastas tikėtinas tūris be vienetų: {val}")
                 return val
            # Standartiniai butelių dydžiai (ml) -> konvertuojame
            if val in [187, 200, 375, 500, 700, 750, 1000, 1500]:
                 logging.info(f"Rastas tikėtinas tūris ml be vienetų: {val}")
                 return val / 1000.0
        except:
            continue
            
    return None

def extract_abv_from_text(text: str) -> Optional[float]:
    """Bando ištraukti ABV iš teksto eilutės."""
    if not text:
        return None
    
    # Ieškome skaičiaus su % (pvz. "46%", "54,2 %", "12.5%")
    # Naudojame regex, kuris pagauna kablelius ir taškus
    match = re.search(r'(\d+(?:[.,]\d+)?)\s*%', text)
    if match:
        try:
            val_str = match.group(1).replace(',', '.')
            val = float(val_str)
            # Validacija: ABV turi būti protingose ribose (0.5 - 99)
            if 0.5 <= val <= 99.0:
                return val
        except ValueError:
            pass
    return None

def recover_volume_from_surrounding_text(document, page_index, y_center):
    """Bando atkurti tūrį iš aplinkinio teksto, jei jis nebuvo rastas kaip entitetas."""
    # Naudojame extract_all_text_in_range, kad gautume VISAS eilutes aplinkui (pvz. viršutinę eilutę su "1 liter")
    # Tolerancija 0.08 (apie 8% puslapio aukščio) turėtų pagauti eilutę viršuje/apačioje
    text = extract_all_text_in_range(document, page_index, y_center, tolerance=0.08)
    if not text:
        return None
    
    logging.info(f"Ieškoma tūrio tekste (Y={y_center:.4f}): '{text}'")
    return extract_volume_from_text(text)

def recover_abv_from_surrounding_text(document, page_index, y_center):
    """Bando atkurti ABV iš aplinkinio teksto."""
    # Naudojame extract_all_text_in_range
    # Tolerancija 0.04 (4% puslapio aukščio) turėtų pagauti ABV net jei jis šiek tiek pasislinkęs
    text = extract_all_text_in_range(document, page_index, y_center, tolerance=0.04)
    if not text:
        return None
    
    logging.info(f"Ieškoma ABV tekste (Y={y_center:.4f}): '{text}'")
    return extract_abv_from_text(text)
